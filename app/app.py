import requests
import pandas as pd
from sqlalchemy import create_engine
import os
import argparse
import matplotlib.pyplot as plt
from pathlib import Path
from docx import Document

    
def check_updates(data, engine):
    # Convertir API en DataFrame
    df_api = pd.DataFrame(data["results"])
    df_api["fecha_carg"] = pd.to_datetime(df_api["fecha_carg"])

    # bajar datos de la bd
    query = "SELECT fecha_carg as latest_fecha FROM cleaned_df"
    latest_db = pd.read_sql(query, engine)
    latest_db_fecha = pd.to_datetime(latest_db["latest_fecha"].iloc[0,0])

    # Obtener la fecha max
    latest_api_fecha = df_api["fecha_carg"].max()

    # check if the data in the API has changed
    if latest_api_fecha.equals(latest_db_fecha) != True:
        print("Nuevos datos guardados")
        # get latest df
        last_df = df_api[df_api["fecha_carg"] > latest_db_fecha]
        registro_path = "./output/historico/registro_historico.csv"
        registro_df = pd.read_csv(registro_path)
        registro_df = latest_db.to_sql('clean_df',con = engine,if_exists = 'append',index = False)
        registro_df.to_csv(registro_path, index=False)
        save_raw_csv(last_df.to_dict(orient="records"), engine)
    else:
        print("No hay nuevos records")
        return None

def save_raw_csv(data, engine):
    # conseguir datos
    df = pd.DataFrame(data["results"])
    cleaned_df = pd.DataFrame(df[["objectid", "nombre", "direccion", "tipozona", "no2", "pm10", "pm25", "tipoemisio", "fecha_carg", "calidad_am", "fiwareid"]])
    cleaned_df = cleaned_df[cleaned_df.objectid != 22]
    cleaned_df["fecha_carg"] = pd.to_datetime(cleaned_df["fecha_carg"], format="%Y-%m-%dT%H:%M:%S%z")

    # Save in PostgreSQL
    cleaned_df.to_sql("cleaned_df", engine, if_exists="replace", index=False)
    print("Datos cargados a PostgreSQL")
    
    # Guardar el CSV
    os.makedirs("./data/raw", exist_ok=True)
    latest_date = cleaned_df["fecha_carg"].max()
    cleaned_df["fecha_carg"] = cleaned_df["fecha_carg"].dt.strftime("%Y-%m-%dT%H:%M:%S%z")
    safe_timestamp = latest_date.strftime("%Y-%m-%dT%H-%M-%S")
    csv_filepath = os.path.join("./data/raw", f"ultimo_{safe_timestamp}.csv")

    try:
            cleaned_df.to_csv(csv_filepath, index=False)
            print(f"CSV guardado en: {csv_filepath}")
    except:
            print(f"Error al guardar el archivo CSV")

    # get path for registro
    file = "./output/historico/registro_historico.csv"
    if (os.path.exists(file) and os.path.isfile(file)):
        os.remove(file)
        registro_path = "./output/historico/registro_historico.csv"
        registro_df = pd.read_sql("SELECT * FROM cleaned_df", engine)
        registro_df.to_csv(registro_path, index=False)

def parse_args():
    p = argparse.ArgumentParser(description="Reportes")
    p.add_argument("--modo", choices=["actual", "historico"], required=True, help="Modo: 'actual' o 'historico'")
    p.add_argument("--since", type=str, help="Fecha de inicio (formato: YYYY-MM-DD)")
    p.add_argument("--estacion", type=str, help="ID de estacion")
    return p.parse_args()

def generate_actual_report():
    # Get latest csv
    csv_folder = "./data/raw/"
    file = os.listdir(csv_folder)[0]
    if not file:
        print("No hay CSV")
    else:
        latest_csv = os.path.join(csv_folder, file)
        data = pd.read_csv(latest_csv)

        # Create summary tables
        summary = data.groupby("fiwareid")[["no2", "pm10", "pm25"]].agg("mean")
        summary["estacion"] = data.groupby("fiwareid")["nombre"].first()
        os.makedirs("./output/actual", exist_ok=True)
        print("Tabla de resumen guardada")

        # Generate graph
        summary["no2"].plot(kind="bar")
        plt.title("NO2 por estacion")
        plt.xlabel("Estacion")
        plt.ylabel("Nivel NO2") 

        summary["pm10"].plot(kind="bar")
        plt.title("pm10 por estacion")
        plt.xlabel("Estacion")
        plt.ylabel("Nivel pm10")

        summary["pm25"].plot(kind="bar")
        plt.title("pm25 por estacion")
        plt.xlabel("Estacion")
        plt.ylabel("Nivel pm25")

        # Export them
        output_dir = os.path.join("./output/actual")
        os.makedirs(output_dir, exist_ok=True)

        summary.to_csv(os.path.join(output_dir, "tabla_actual.csv"))

        plt.savefig(os.path.join(output_dir, "no2_por_estacion.png"))
        plt.savefig(os.path.join(output_dir, "pm10_por_estacion.png"))
        plt.savefig(os.path.join(output_dir, "pm25_por_estacion.png"))
        doc = Document()
        # Add a title
        doc.add_heading('Grafica informe actual estacion/no2', level=1)
        # Add a paragraph
        doc.add_paragraph('Informe actual sobre el No2 generado según la estación')
        doc.add_picture("./output/actual/no2_por_estacion.png")
        doc.add_heading('Grafica informe actual estacion/pm10', level=1)
        # Add a paragraph
        doc.add_paragraph('Informe actual sobre el pm10 generado según la estación')
        # Add the saved plot image to the document
        doc.add_picture("./output/actual/pm10_por_estacion.png" )
        # Save the document
        doc.add_heading('Grafica informe actual estacion/pm25', level=1)
        # Add a paragraph
        doc.add_paragraph('Informe actual sobre el pm25 generado según la estación')
        # Add the saved plot image to the document
        doc.add_picture("./output/actual/pm25_por_estacion.png")

        save_path = os.path.join(output_dir, "Reporte_Actual.docx")
        doc.save(save_path)
        # Save the document

        print("Graficos actuales guardados")

def generate_historico_report(since, estacion):
    # Get latest csv
    historico_df = pd.read_csv("./output/historico/registro_historico.csv")
    # use filters with the given info
    if since:
        since_date = historico_df[pd.to_datetime(since)]
        historico_df = historico_df[historico_df["fecha_carg"] >= since_date]
    if estacion:
        historico_df = historico_df[historico_df["fiwareid"] == estacion]

    for station_id in historico_df["fiwareid"].unique():
        station_data = historico_df[historico_df["fiwareid"] == station_id]

        plt.plot(pd.to_datetime(station_data["fecha_carg"]), station_data["no2"], label=str(station_id))

    # generate grap
    plt.title("NO2 histórico por estación")
    plt.xlabel("Fecha")
    plt.ylabel("Nivel NO2")
    plt.legend()
    plt.tight_layout()

    # save it 
    output_dir = os.path.join(os.getcwd(), "output", "historico")
    os.makedirs(output_dir, exist_ok=True)
    plt.savefig(os.path.join(output_dir, "NO2_historico.png"))
    plt.close()
    print("Graficos historicos guardados")

def main():
    args = parse_args()

    # Get data from API
    url = "https://valencia.opendatasoft.com/api/explore/v2.1/catalog/datasets/estacions-contaminacio-atmosferiques-estaciones-contaminacion-atmosfericas/records"
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        data = response.json()  # dict with keys like 'results', 'total_count'
        print("Conexión exitosa a la API")
    except:
        print("Error en la conexión")
        return

    # Connect to PostgreSQL
    churro = 'postgresql://postgres:mysecretpassword@localhost:5432/postgres'
    engine = create_engine(churro)


    save_raw_csv(data, engine)
    
    if args.modo == "historico":
        generate_historico_report(args.since, args.estacion)
    elif args.modo == "actual":
        generate_actual_report()

if __name__ == "__main__":
    main()

# Ingesta + informe actual
# python app/app.py --modo actual

# Ingesta + informe histórico
# python app/app.py --modo historico

# Opcionales útiles
# python app/app.py --since "2025-10-01" --estacion "A05_POLITECNIC_60m"